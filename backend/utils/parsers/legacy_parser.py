from typing import List, Dict, Any, Optional
import logging
import io
import PyPDF2
import docx

from .base import DocumentParser, ParsedDocument, ParseError

logger = logging.getLogger(__name__)


class LegacyParser(DocumentParser):
    """
    Legacy document parser using PyPDF2 and python-docx for basic document processing.
    Used as fallback when Docling is not available or fails.
    
    Features:
    - Basic PDF text extraction
    - DOCX document processing
    - Plain text and markdown support
    - Simple metadata extraction
    """
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """
        Initialize legacy parser with configuration.
        
        Args:
            config: Configuration dictionary for parser settings
        """
        self.config = config or {}
        self.supported_types = {
            "application/pdf": self._process_pdf,
            "text/plain": self._process_text,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": self._process_docx,
            "text/markdown": self._process_markdown,
        }
        
        logger.info("Legacy parser initialized with %d supported types", len(self.supported_types))
    
    def can_parse(self, file_type: str, filename: str) -> bool:
        """
        Check if legacy parser can handle this file type
        
        Args:
            file_type: MIME type of the file
            filename: Name of the file
            
        Returns:
            bool: True if parser can handle this file type
        """
        return file_type in self.supported_types
    
    def get_supported_types(self) -> List[str]:
        """
        Get list of MIME types supported by legacy parser
        
        Returns:
            List[str]: List of supported MIME types
        """
        return list(self.supported_types.keys())
    
    async def parse(self, content: bytes, filename: str) -> ParsedDocument:
        """
        Parse document using legacy methods
        
        Args:
            content: File content as bytes
            filename: Name of the file
            
        Returns:
            ParsedDocument: Parsed document with basic structure
        """
        try:
            # Determine file type - this should be passed from caller
            file_type = self._detect_file_type(filename)
            
            if file_type in self.supported_types:
                processor_method = self.supported_types[file_type]
                text_content = await processor_method(content)
            else:
                # Fallback to treating as text
                text_content = content.decode("utf-8", errors="ignore")
            
            # Clean and prepare content
            cleaned_content = self._clean_text(text_content)
            
            # Create metadata
            metadata = {
                'filename': filename,
                'file_size': len(content),
                'parser_type': 'legacy',
                'word_count': len(cleaned_content.split()),
                'character_count': len(cleaned_content)
            }
            
            return ParsedDocument(
                content=cleaned_content,
                metadata=metadata,
                structure=None,  # Legacy parser doesn't extract structure
                tables=None,     # Legacy parser doesn't extract tables
                images=None      # Legacy parser doesn't extract images
            )
            
        except Exception as e:
            logger.error(f"Legacy parsing failed for {filename}: {str(e)}")
            raise ParseError(
                f"Failed to parse document with legacy parser: {str(e)}", 
                "LegacyParser", 
                e
            )
    
    def _detect_file_type(self, filename: str) -> str:
        """Detect file type from filename extension"""
        extension_mapping = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.txt': 'text/plain',
            '.md': 'text/markdown'
        }
        
        file_ext = filename.lower().split('.')[-1]
        return extension_mapping.get(f'.{file_ext}', 'text/plain')
    
    async def _process_pdf(self, content: bytes) -> str:
        """Extract text from PDF content using PyPDF2"""
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return text
            
        except Exception as e:
            logger.error(f"PDF processing error: {str(e)}")
            raise ParseError(f"Failed to process PDF: {str(e)}", "LegacyParser", e)
    
    async def _process_docx(self, content: bytes) -> str:
        """Extract text from DOCX content using python-docx"""
        try:
            doc_file = io.BytesIO(content)
            doc = docx.Document(doc_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return text
            
        except Exception as e:
            logger.error(f"DOCX processing error: {str(e)}")
            raise ParseError(f"Failed to process DOCX: {str(e)}", "LegacyParser", e)
    
    async def _process_text(self, content: bytes) -> str:
        """Process plain text content"""
        try:
            return content.decode("utf-8", errors="ignore")
        except Exception as e:
            logger.error(f"Text processing error: {str(e)}")
            raise ParseError(f"Failed to process text: {str(e)}", "LegacyParser", e)
    
    async def _process_markdown(self, content: bytes) -> str:
        """Process Markdown content"""
        try:
            return content.decode("utf-8", errors="ignore")
        except Exception as e:
            logger.error(f"Markdown processing error: {str(e)}")
            raise ParseError(f"Failed to process markdown: {str(e)}", "LegacyParser", e)
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text content"""
        # Remove excessive whitespace
        lines = text.split("\n")
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if line:  # Skip empty lines
                cleaned_lines.append(line)
        
        # Join with single newlines
        cleaned_text = "\n".join(cleaned_lines)
        
        # Remove excessive spaces
        import re
        cleaned_text = re.sub(r" +", " ", cleaned_text)
        
        return cleaned_text
